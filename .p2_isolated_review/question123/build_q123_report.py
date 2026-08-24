#!/usr/bin/env python3
"""生成不覆盖旧文档的《C题问题1—3统合报告_v1》草稿。"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

PROJECT = Path(__file__).resolve().parents[1]
DOC = PROJECT / "doc"
OUT = DOC / "C题问题1-3统合报告_v1_pre_equations.docx"
FINAL = DOC / "C题问题1-3统合报告_v1.docx"
FIG = PROJECT / "figures" / "q123"
Q1 = PROJECT / "results" / "q1"
Q2 = PROJECT / "results" / "q2"
Q3 = PROJECT / "results" / "q3"

BLUE = "1F4E79"
LIGHT = "D9EAF7"
GRAY = "F2F2F2"
INK = "202124"
PRESET_NAME = "compact_reference_guide"
# 命名覆盖 cumcm_a4_math_paper：竞赛中文技术报告使用A4、宋体正文与居中封面；
# 其余正文节奏、蓝色层级、9360-DXA表格遵循 compact_reference_guide。
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
TABLE_COUNTER = 0


def set_font(run, size=10.5, bold=False, color=INK, east="宋体"):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def repeat_header(row):
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(node)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText"); text.set(qn("xml:space"), "preserve"); text.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, end])


def body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def equation(doc, placeholder):
    p = doc.add_paragraph(placeholder)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    set_font(p.runs[0], 11)


def table(doc, headers, rows, widths=None, font=8.5, caption="关键结果"):
    global TABLE_COUNTER
    TABLE_COUNTER += 1
    cap = doc.add_paragraph(f"表{TABLE_COUNTER} {caption}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = True
    for r in cap.runs: set_font(r, 8.5, color="555555")
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_header(t.rows[0])
    for i, value in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = str(value); shade(c, LIGHT)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c.paragraphs[0].runs: set_font(r, font, True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            for r in cells[i].paragraphs[0].runs: set_font(r, font)
    if widths:
        total = float(sum(widths))
        dxa = [round(TABLE_WIDTH_DXA * float(w) / total) for w in widths]
        dxa[-1] += TABLE_WIDTH_DXA - sum(dxa)
        t.autofit = False
        tbl_pr = t._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA)); tbl_w.set(qn("w:type"), "dxa")
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            inserted = False
            for tag in ("w:tblBorders", "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook"):
                node = tbl_pr.find(qn(tag))
                if node is not None:
                    node.addprevious(tbl_ind); inserted = True; break
            if not inserted: tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA)); tbl_ind.set(qn("w:type"), "dxa")
        grid = t._tbl.tblGrid
        for child in list(grid): grid.remove(child)
        for value in dxa:
            col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(value)); grid.append(col)
        for row in t.rows:
            for c, value in zip(row.cells, dxa):
                tc_w = c._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW"); c._tc.get_or_add_tcPr().append(tc_w)
                tc_w.set(qn("w:w"), str(value)); tc_w.set(qn("w:type"), "dxa")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def figure(doc, filename, caption, alt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # 图4为横向直方图，略缩宽度可避免正文与图被拆到两页。
    width = 5.20 if filename == "raw_q2_bmi_distribution.png" else 5.95
    shape = run.add_picture(str(FIG / filename), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    for r in cap.runs: set_font(r, 8.5, color="555555")


def setup(doc):
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.25)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for name, size, color in (("Title", 22, BLUE), ("Heading 1", 15, BLUE), ("Heading 2", 12, BLUE)):
        st = styles[name]; st.font.name = "Times New Roman"; st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color)
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_page_number(footer)
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None: zoom.set(qn("w:percent"), "100")


def main():
    DOC.mkdir(parents=True, exist_ok=True)
    coef = pd.read_csv(Q1 / "q1_model_coefficients.csv")
    q2 = pd.read_csv(Q2 / "问题2核心方案对比.csv")
    q2_rec = pd.read_csv(Q2 / "问题2最终建议表.csv")
    q3_rec = pd.read_csv(Q3 / "问题3最终建议表_v2.csv")
    q3_nested = json.loads((Q3 / "问题3嵌套外层独立审计.json").read_text(encoding="utf-8"))
    q3_audit = pd.read_csv(Q3 / "最终政策外层4折审计.csv")
    model_cmp = pd.read_csv(Q3 / "模型外层测试对比.csv")

    doc = Document(); setup(doc)
    title = doc.add_paragraph("NIPT检测时点与BMI分组优化\n——C题问题1—3精简统合报告", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("建模草稿 / 内部验证版本 v1（不替代临床诊断）")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs: set_font(r, 11, color="666666")
    doc.add_paragraph("\n")
    table(doc, ["问题", "最终回答", "证据边界"], [
        ["问题1", "孕周与Y浓度正相关，BMI与其负相关；交互不显著", "纵向关联，不作因果解释"],
        ["问题2", "BMI<34.357：18周；BMI≥34.357：22周", "题面三级风险下的内部主推荐"],
        ["问题3", "简化为BMI<35：18周；BMI≥35：22周", "开发后固定政策的内部稳定性建议"],
    ], [2, 8, 6], 9, "问题1—3最终回答与证据边界")
    body(doc, "由表1先给出三问的最终回答；后文分别给出模型、验证与限制。")
    body(doc, "说明：本文件统合问题1—3的最终可用证据，旧搜索树、局部贪心和多组探索仅在方法演进中简述。所有达标概率均指Y染色体浓度达到题面4%阈值的概率，不是胎儿染色体异常诊断准确率。")
    doc.add_page_break()

    heading(doc, "摘 要", 1)
    body(doc, "针对267名男胎孕妇的1082条重复检测记录，本文先以随机截距线性混合模型分析Y染色体浓度与检测孕周、BMI的条件关联，再把首次达到4%的孕周构造成患者级删失结局，采用区间删失XGBoost-AFT生成孕周达标概率，并在连续BMI候选空间内优化分组和推荐时点。问题1显示孕周线性项和二次项为正、BMI项为负，孕周×BMI交互不显著。问题2在题面三级风险口径下推荐以BMI 34.357分组并分别于18周、22周检测；16/18仅作为中期内部等待损失连续增加时的敏感性方案。问题3显示多因素相对仅BMI模型的增量很小；无外层选参的嵌套审计为均值安全4/4折、四分位审计层LCB安全3/4折。综合可解释性后给出整数切点35与18/22周的开发后固定政策，其稳定性证据来自内部回放和Bootstrap，并非完全独立的无偏测试。")
    p = doc.add_paragraph(); set_font(p.add_run("关键词："), bold=True); set_font(p.add_run("NIPT；胎儿浓度；线性混合模型；区间删失；AFT；Optimal Binning；嵌套交叉验证"))

    heading(doc, "1 问题重述与数据层级", 1)
    body(doc, "题目要求依次回答Y浓度与孕周、BMI等因素的关系，按BMI分组给出最佳NIPT时点，并在综合多因素和检测误差后修订分组。已有研究指出胎儿游离DNA浓度与孕周、母体体重等因素相关[1,2]。附件包含男胎1082条记录、267名孕妇；同一孕妇可能重复检测。因此问题1保留记录级纵向结构，问题2、3以孕妇为单位构造达标时间和验证折。")
    table(doc, ["层级", "样本单位", "用途", "防泄漏规则"], [
        ["问题1", "1082条记录/267人", "Y浓度关联模型", "孕妇随机截距"],
        ["问题2、3", "267名孕妇", "达标时间和时点决策", "同一孕妇不跨折；患者级Bootstrap"],
    ], [2.5, 3.5, 5, 5], caption="三个子问题的数据层级")
    body(doc, "由表2可见三问共享原始附件，但分析单位和防泄漏措施不同；原始记录结构见图1。")
    figure(doc, "raw_q1_y_week_bmi.png", "图1 记录级孕周、BMI与Y染色体浓度", "散点图展示检测孕周、Y浓度和记录级BMI，并标出4%达标线。")

    heading(doc, "2 数据预处理与统一假设", 1)
    body(doc, "孕周w+d统一换算为w+d/7周，Y浓度按比例计算，题面4%写为0.04。问题2、3的BMI、年龄、身高、体重、IVF、孕次和产次均取同一孕妇记录中位数；缺失则终止，不做隐式填补。首次观测即达标记为左删失，前次未达标而后次达标记为区间删失，最后一次仍未达标记为右删失。")
    body(doc, "核心假设包括：附件记录可代表本样本内的浓度变化；在10—25周候选窗内，AFT条件分布可用于比较政策；患者级重采样足以描述内部抽样波动。题面把≤12周、13—27周和≥28周分为三级风险，因此18周和22周在题面上同属中期。")

    heading(doc, "3 问题1：纵向关联模型", 1)
    heading(doc, "3.1 模型与检验", 2)
    body(doc, "令Yij为孕妇i第j次检测的Y浓度，Wij、Bij为该记录标准化孕周和BMI，ui为孕妇随机截距。随机效应模型适合处理纵向重复观测相关性[5]。主模型为：")
    equation(doc, "{{EQ1}}")
    body(doc, "二次项刻画孕周曲率，交互项检验BMI是否改变孕周斜率。扩展模型增加年龄、身高、IVF、唯一比对读段数、比对比例和GC含量；基础与扩展模型用最大似然拟合，单项采用Wald检验，整体增量采用似然比检验。图2给出残差与正态性诊断。")
    figure(doc, "process_q1_diagnostics.png", "图2 问题1基础模型残差诊断", "左为残差对拟合值，右为残差Q-Q图。")

    heading(doc, "3.2 结果", 2)
    base = coef[(coef["模型"] == "基础模型") & (coef["变量"] != "Intercept")]
    label = {"week_c":"孕周", "week_c2":"孕周²", "bmi_c":"BMI", "week_c:bmi_c":"孕周×BMI"}
    rows = [[label[r["变量"]], f'{r["系数"]:.4f}', f'{r["标准误"]:.4f}', f'{r["p值"]:.3g}', f'[{r["95%CI下限"]:.4f},{r["95%CI上限"]:.4f}]'] for _, r in base.iterrows()]
    table(doc, ["中心化条件效应", "系数", "标准误", "p值", "95%CI"], rows, [4, 2.5, 2.5, 2.5, 4.5], caption="问题1基础混合模型固定效应")
    body(doc, "由表3可见：孕周线性项0.1668（p<10⁻⁴⁷）、二次项0.0428（p=5.80×10⁻⁶）均为正；BMI项−0.0719（p=0.00250）为负；交互项p=0.336，不支持不同BMI具有显著不同的孕周增长斜率。扩展模型AIC由836.330降至817.693，LRT=30.637、df=6、p=2.97×10⁻⁵，说明辅助变量整体增加统计解释量，但不等于因果机制；图3显示主要效应区间。")
    figure(doc, "result_q1_effects.png", "图3 问题1基础混合模型效应及95%置信区间", "森林图展示标准化中心化条件效应；零线表示无效应。")

    heading(doc, "4 问题2：BMI分组与最佳时点", 1)
    heading(doc, "4.1 区间删失AFT与全局分箱", 2)
    body(doc, "设Ti为首次达到4%的孕周，Fi(t)=P(Ti≤t|xi)。左删失、区间删失和右删失的似然贡献分别如下，训练最小化负对数似然。XGBoost-AFT可直接接收下、上界，并在外4折、内3折患者级验证中完成早停与概率预测[3,7]。嵌套验证用于降低调参与误差估计复用造成的偏差[4]。")
    equation(doc, "{{EQ2}}")
    body(doc, "在冻结的训练OOF概率上，按BMI排序并在相邻不同BMI值之间进行全局Optimal Binning搜索；候选区间连续，每组至少30人。该方法在给定候选空间内全局比较分段，不等同于局部贪心决策树；图4显示切点位于样本分布的高BMI尾部。")
    figure(doc, "raw_q2_bmi_distribution.png", "图4 问题2BMI分布与Optimal Binning切点", "患者级BMI直方图；竖线为34.357。")

    heading(doc, "4.2 三方案判优与主推荐", 2)
    rows = []
    for _, r in q2.iterrows():
        rows.append([r["方案"], f'{r["患者平均孕周"]:.3f}', f'{r["总体预测达标概率"]:.4f}', f'{r["Bootstrap安全率"]:.4f}', f'{r["重复折安全率_y"]:.2f}', f'{r["状态改变率均值"]:.4f}'])
    table(doc, ["方案", "平均周", "总体达标概率", "Bootstrap安全率", "重复折安全率", "状态改变率"], rows, [2.5,2.5,3,3,3,3], caption="问题2三套固定时点方案比较")
    body(doc, "由表4给出三套离散方案。问题2固定政策的部署安全线为组内最低概率0.89；稳定门槛为Bootstrap安全率≥0.95、重复折安全率≥0.90、100次重复四折中4/4折全部安全的比例≥0.70。18/22三项稳定性均为1.0，故在题面三级风险下作为主推荐：BMI<34.357的215人建议18周，BMI≥34.357的52人建议22周。图5展示三类稳定性指标，图6展示总体达标概率与平均推荐周的权衡。")
    figure(doc, "process_q2_policy_stability.png", "图5 问题2三套离散政策的稳定性点图", "分别展示Bootstrap安全率、重复折安全率和重复4/4比例；离散政策之间不连线。")
    figure(doc, "result_q2_tradeoff.png", "图6 问题2总体达标概率与平均推荐孕周权衡", "三个离散方案的总体均值；不与组内最低安全线混用。")

    heading(doc, "4.3 连续等待风险敏感性", 2)
    body(doc, "为回答“是否越早越好”，扩展分析定义准确性损失Ra、连续等待损失Rt和不稳定损失Rs：")
    equation(doc, "{{EQ5}}")
    body(doc, "三类损失组成加权风险：")
    equation(doc, "{{EQ3}}")
    body(doc, "当ws=0时，16/18仅在0.8022≤wa≤0.8461（0.1539≤wt≤0.1978）范围内最优。因此16/18是中期内部等待损失连续增加假设下的折中方案，不与题面18/22主答案并列。")

    heading(doc, "5 问题3：多因素约束下的稳健分组", 1)
    heading(doc, "5.1 多因素是否值得增加复杂度", 2)
    means = model_cmp.groupby("模型")["外层AFT负对数似然"].mean().to_dict()
    body(doc, f"多因素AFT和仅BMI AFT的外层负对数似然均值分别为{means.get('多因素', float('nan')):.6f}和{means.get('仅BMI', float('nan')):.6f}，差异约−0.00101；配对区间评分差的Bootstrap区间跨0（−0.01880至0.02885）。因此辅助变量可用于风险校准，但不足以支持大量分组或逐人政策。患者级原始变量中身高、体重和BMI存在相关结构，置换重要性只作预测证据，不作因果解释；图7给出原始协变量相关矩阵。")
    figure(doc, "raw_q3_feature_evidence.png", "图7 问题3患者级原始协变量相关矩阵", "由男胎附件按孕妇中位数聚合；颜色表示Pearson相关系数。")

    heading(doc, "5.2 无外层选参的嵌套选组审计", 2)
    body(doc, "开发可靠性q=0.95预先固定。BMI稳定排序后仅在不同值之间切分，每组至少30人；区间g在周t的保守下限及动态规划目标为：")
    equation(doc, "{{EQ4}}")
    body(doc, f"每个外层折内只依据开发平均孕周判优，并在0.25周近优范围内选最少组；外层指标不参与选择。图8所示结果为审计层均值安全{q3_nested['mean_safe_folds']}/4折、LCB安全{q3_nested['lcb_safe_folds']}/4折，平均测试孕周{q3_nested['mean_test_week']:.3f}，平均测试达标概率{q3_nested['mean_test_readiness']:.4f}；四折分别选1组两次、2组两次。这一结果验证的是选组流程，不是最终切点35政策的无偏测试。")
    figure(doc, "process_q3_group_selection.png", "图8 问题3无外层选参的逐折审计", "左图标注每折选中组数；右图给出四分位审计层最低均值与LCB。")

    heading(doc, "5.3 开发后固定政策与建议表", 2)
    table(doc, ["BMI分组", "人数", "推荐时点", "性质"], q3_rec.astype(str).values.tolist(), [4,2.5,3,6.5], caption="问题3开发后固定政策建议")
    body(doc, "由表5给出最终简化建议：BMI<35的227人建议18周，BMI≥35的40人建议22周。固定政策内部回放的四个外层折中，政策组均值和LCB均4/4通过，5000次患者Bootstrap安全率为1.0；总体达标概率约0.9533。切点Bootstrap 95%分位区间为[30.494,36.158]，两组重复中18/23出现449次而18/22仅80次，故35与18/22应表述为开发完成后的简化部署政策，而非Bootstrap众数或独立无偏泛化结果；图9只表示固定政策内部回放。")
    figure(doc, "result_q3_final_policy.png", "图9 问题3整数BMI政策及固定政策内部回放", "左为BMI 35分组和18/22周；右为开发后固定政策的内部回放。")

    heading(doc, "6 测量误差与稳健性", 1)
    table(doc, ["项目", "问题2主政策", "问题3固定政策", "解释"], [
        ["Y浓度误差SD", "0.003716", "沿用同源估计", "约0.372个百分点"],
        ["状态改变率", "0.027154", "0.027236", "扰动后4%删失状态变化"],
        ["Bootstrap安全率", "1.000", "1.000", "患者级内部重采样"],
    ], [3.5,3.5,3.5,6], caption="问题2与问题3的测量误差和重采样结果")
    body(doc, "由表6汇总误差与重采样结果。误差模拟检查4%阈值附近的状态扰动，不等同于胎儿疾病误诊率。患者Bootstrap以孕妇为抽样单元[6]，避免把同一孕妇多条记录当作独立样本。")

    heading(doc, "7 模型评价、优点与局限", 1)
    body(doc, "评价顺序为：先满足无泄漏和可靠性约束，再比较稳定性与平均推荐孕周，最后考虑模型复杂度与解释性。优点是保留纵向重复测量、显式处理三类删失、以患者为验证单位，并把题面风险与连续风险敏感性分开。")
    body(doc, "局限性包括：样本仅267名孕妇且来自单一数据源；高BMI尾部人数较少；切点区间较宽；LCB是预设保守评分而非严格临床置信保证；多因素增益不稳定；最终35与18/22经过全数据探索，尚无完全独立测试集或外部医院验证。因此结果只能用于数学建模与研究性检测时点分流，不能作为临床诊断或处置依据。")

    heading(doc, "8 结论", 1)
    table(doc, ["问题", "结论", "证据等级/适用范围"], [
        ["1", "孕周与Y浓度正相关，BMI与其负相关；交互不显著", "条件关联；内部纵向模型"],
        ["2", "BMI<34.357于18周，BMI≥34.357于22周", "题面三级风险的内部主推荐"],
        ["2敏感性", "BMI<34.357于16周，BMI≥34.357于18周", "仅在特定连续风险权重区间"],
        ["3", "BMI<35于18周，BMI≥35于22周", "开发后固定政策内部稳定性建议"],
    ], [2.2,8.5,6], caption="问题1—3最终结论与适用范围")
    body(doc, "由表7汇总最终结论。综合而言，更多分组和更多变量没有带来足够稳定的净收益。当前最合理的工程方案是保留问题2精确切点作为算法结果，在问题3用整数35形成便于解释的两组政策，同时明确其内部验证边界。")

    heading(doc, "参考文献", 1)
    refs = [
        "[1] Ashoor G, Syngelaki A, Poon LCY, Rezende JC, Nicolaides KH. Fetal fraction in maternal plasma cell-free DNA at 11–13 weeks’ gestation: relation to maternal and fetal characteristics. Ultrasound in Obstetrics & Gynecology, 2013, 41(1): 26–32. DOI:10.1002/uog.12331.",
        "[2] Wang E, Batey A, Struble C, Musci T, Song K, Oliphant A. Gestational age and maternal weight effects on fetal cell-free DNA in maternal plasma. Prenatal Diagnosis, 2013, 33(7): 662–666. DOI:10.1002/pd.4119.",
        "[3] Barnwal A, Cho H, Hocking TD. Survival Regression with Accelerated Failure Time Model in XGBoost. Journal of Computational and Graphical Statistics, 2022, 31(4): 1292–1302. DOI:10.1080/10618600.2022.2067548.",
        "[4] Varma S, Simon R. Bias in error estimation when using cross-validation for model selection. BMC Bioinformatics, 2006, 7:91. DOI:10.1186/1471-2105-7-91.",
        "[5] Laird NM, Ware JH. Random-effects models for longitudinal data. Biometrics, 1982, 38(4): 963–974. DOI:10.2307/2529876.",
        "[6] Efron B, Tibshirani RJ. An Introduction to the Bootstrap. New York: Chapman & Hall, 1993.",
        "[7] Chen T, Guestrin C. XGBoost: A Scalable Tree Boosting System. Proceedings of KDD 2016: 785–794. DOI:10.1145/2939672.2939785.",
    ]
    for ref in refs:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.7); p.paragraph_format.first_line_indent = Cm(-0.7); p.paragraph_format.space_after = Pt(3); set_font(p.add_run(ref), 9)

    heading(doc, "附录A 复现与文件范围", 1)
    manifest = json.loads((PROJECT / "results" / "q123" / "复现清单.json").read_text(encoding="utf-8"))
    manifest_items = sum(len(manifest[key]) for key in ("code", "results", "figures"))
    body(doc, f"唯一轻量入口为 E:\\anaconda\\python.exe C题\\question123\\run_q123.py。它重建问题1、校验问题2冻结政策、重建问题3无外层选参审计、校验问题2/3冻结结果哈希、重绘9组证据图并生成{manifest_items}项复现清单。历史大规模搜索不在默认命令中重跑。附件SHA-256为14827156218bd4f7e4f16db4aa6d9f757c6648379e038ae6c6b58383648614af。")
    body(doc, "本报告为建模草稿与内部分析参考；正式参赛前仍需结合官方当年格式、队伍信息和篇幅要求二次排版。")

    doc.save(OUT)
    print(f"draft={OUT}")


if __name__ == "__main__":
    main()
