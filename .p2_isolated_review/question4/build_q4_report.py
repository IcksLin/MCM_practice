"""构建C题问题4独立技术报告（新版本，不覆盖旧报告）。"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

import pandas as pd
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
FIGURES = ROOT / "figures"
DOC_DIR = ROOT / "doc"
OUTPUT = DOC_DIR / "C题问题4最终报告_v5.docx"

DOCX_SCRIPTS = Path(r"C:\Users\admin\.codex\skills\math-modeling\tools\docx\scripts")
sys.path.insert(0, str(DOCX_SCRIPTS))
import paper_format as pf


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t"); text.text = "1"
    run.append(text); fld.append(run); paragraph._p.append(fld)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def style_table(table, widths=None):
    table.autofit = False
    set_repeat_header(table.rows[0])
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths and c < len(widths):
                cell.width = Cm(widths[c])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    pf.set_run_font(run, "宋体", 9, r == 0)


def add_table(doc, rows, caption, widths=None):
    pf.body(doc, f"如{caption.split(' ')[0]}所示，表内数值均来自冻结后的外层留出预测或患者级Bootstrap，不使用训练拟合值冒充泛化结果。")
    pf.figure_caption(doc, caption)
    table = pf.three_line_table(doc, rows)
    style_table(table, widths)
    return table


def add_figure(doc, filename, caption, discussion, width=15.2):
    number = caption.split()[0].replace("图", "")
    pf.body(doc, f"参见图{number}。{discussion}")
    image_paragraph = pf.image(doc, FIGURES / filename, width_cm=width)
    # 为 Word 中的科研图补充无障碍替代文本；文字同时说明图号与图意。
    for doc_pr in image_paragraph._p.xpath(".//wp:docPr"):
        doc_pr.set("title", caption)
        doc_pr.set("descr", f"{caption}。{discussion}")
    pf.figure_caption(doc, caption)


def add_bullets(doc, items):
    for item in items:
        p = pf.paragraph(doc)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        pf.set_run_font(p.add_run("• "), bold=True)
        pf.set_run_font(p.add_run(item))


def fmt(x, digits=3):
    return f"{float(x):.{digits}f}"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=OUTPUT)
    args = parser.parse_args()
    output = args.output.resolve()
    if output.exists():
        raise FileExistsError(f"报告已存在，拒绝覆盖：{output}")
    output.parent.mkdir(parents=True, exist_ok=True)
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    pred = pd.read_csv(RESULTS / "q4_outer_predictions.csv")
    boot = pd.read_csv(RESULTS / "q4_bootstrap_summary.csv")
    metrics = pd.read_csv(RESULTS / "q4_metrics_summary.csv")
    audit = pd.read_csv(RESULTS / "q4_fold_audit.csv")
    method = json.loads((RESULTS / "q4_final_method.json").read_text(encoding="utf-8"))
    actions = pd.read_csv(RESULTS / "q4_final_recommendations.csv")["recommended_action"].value_counts()

    def b(target, metric):
        row = boot[(boot.target == target) & (boot.metric == metric)].iloc[0]
        return float(row.point), float(row.ci_low), float(row.ci_high)

    any_ap = b("any", "pr_auc"); any_sens = b("any", "sensitivity")
    any_spec = b("any", "specificity"); any_prec = b("any", "precision")
    any_f2 = b("any", "f2"); any_brier = b("any", "brier")

    doc = pf.new_document(contest="cumcm")
    for section in doc.sections:
        section.top_margin = Cm(2.3); section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.6); section.right_margin = Cm(2.4)
        add_page_number(section.footer.paragraphs[0])
    doc.core_properties.title = "C题问题4：女胎T13/T18/T21非整倍体判定"
    doc.core_properties.subject = "患者分组嵌套交叉验证、代价敏感XGBoost与复测分流"
    doc.core_properties.author = "数学建模项目组"
    doc.core_properties.keywords = "NIPT; T13; T18; T21; XGBoost; 嵌套交叉验证"

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70); p.paragraph_format.space_after = Pt(20)
    r = p.add_run("C题问题4：女胎T13/T18/T21\n非整倍体判定")
    pf.set_run_font(r, "黑体", 22, True); r.font.color.rgb = RGBColor(31, 78, 121)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p2.add_run("独立技术报告 · 最终版"); pf.set_run_font(rr, "黑体", 15, False); rr.font.color.rgb = RGBColor(89, 89, 89)
    doc.add_paragraph(); doc.add_paragraph()
    for line in ("数据：附件女胎检测数据，605条记录、147名孕妇",
                 "验证：外层4折/内层3折患者分组嵌套交叉验证",
                 "主方法：代价敏感XGBoost + Platt校准 + 训练折阈值",
                 "定位：问题4独立技术报告；不等同于临床诊断报告"):
        q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.set_run_font(q.add_run(line), "宋体", 11, False)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER; p3.paragraph_format.space_before = Pt(40)
    p3.add_run("生成日期：2026年8月24日").font.size = Pt(11)
    pf.page_break(doc)

    pf.abstract_title(doc)
    pf.body(doc, "本报告解决女胎无创产前检测记录中的T13、T18、T21非整倍体筛查判定问题。数据共有605条重复检测记录，来自147名孕妇；任一异常标签67条，其中T13、T18、T21分别为23、46、13条。由于同一孕妇可多次检测，若按记录随机划分会泄漏个体信息，因此采用外层4折、内层3折的患者分组嵌套交叉验证；全部预处理、超参数选择、概率校准、任一异常来源选择和阈值确定均限制在相应外层训练集内。")
    pf.body(doc, "模型以三个代价敏感XGBoost分类器刻画染色体Z值、GC含量、读段质量、孕周和BMI之间的非线性关系，并以患者等权抑制重复检测次数差异。内层OOF概率经Platt方法校准；判定阈值优先满足灵敏度不低于0.85，再最大化特异度。三个染色体概率的并集与直接任一异常模型在每个外层训练集内比较，四折均选择概率并集。")
    pf.body(doc, f"完全留出的pooled OOF结果显示，任一异常PR-AUC为{any_ap[0]:.3f}，灵敏度为{any_sens[0]:.3f}，特异度为{any_spec[0]:.3f}，精确率为{any_prec[0]:.3f}，F2为{any_f2[0]:.3f}，Brier分数为{any_brier[0]:.3f}。患者级2000次Bootstrap给出的灵敏度95%区间为[{any_sens[1]:.3f}, {any_sens[2]:.3f}]，PR-AUC区间为[{any_ap[1]:.3f}, {any_ap[2]:.3f}]。T18的辨别力相对最好，T13较弱，T21受阳性样本极少影响而不稳定。")
    pf.body(doc, "最终建议采用四级互斥分流：质量异常先复检，任一染色体概率接近阈值时复检，其余记录按任一异常概率给出筛查阳性或筛查阴性。必须强调，附件AE列605条均为“是”，无法提供真实核型或出生结局；本模型只复现AB列筛查标签，不能替代遗传咨询、羊水穿刺等诊断性检查。")
    pf.keywords(doc, "NIPT；非整倍体；患者分组嵌套交叉验证；XGBoost；概率校准；不平衡分类")

    pf.heading1(doc, "1 问题重述与结论边界")
    pf.heading2(doc, "1.1 题目要求")
    pf.body(doc, "第四题要求以女胎记录AB列“染色体的非整倍体”为判定结果，综合X染色体和13、18、21号染色体的Z值、GC含量、读段数及其比例、孕妇BMI等因素，给出女胎异常判定方法。输出不仅要有模型，还要说明如何训练、如何选择阈值、如何验证泛化能力，以及低质量或不确定记录如何处理。")
    pf.heading2(doc, "1.2 目标变量与医学边界")
    pf.body(doc, "AB列允许同一记录同时出现多个标签，因此本质是多标签筛查问题，而不是四选一多分类问题。本报告分别建立T13、T18和T21三个二分类器，再形成“任一异常”总风险。NIPT属于筛查而非确诊；Koumbaris等和Lee等的研究均强调cfDNA计数与质量控制对T13/T18/T21筛查的重要性[1-2]。本题附件没有真实核型金标准，所有性能结论均限定为对AB列的内部复现能力。")
    add_bullets(doc, ["学习目标：复现AB列中的T13、T18、T21筛查标签。",
                      "判定目标：尽量降低漏检，同时控制无意义的过度阳性。",
                      "使用边界：研究性分流工具，不输出“胎儿患病”诊断。",
                      "失败条件：若换用真实临床结局、不同实验平台或不同人群，必须重新训练和外部验证。"])

    pf.heading1(doc, "2 数据固化、清洗与探索")
    pf.heading2(doc, "2.1 数据快照与标签分布")
    pf.body(doc, "输入附件以SHA-256固定为14827156218bd4f7e4f16db4aa6d9f757c6648379e038ae6c6b58383648614af。女胎工作表含605条记录、31列、147名孕妇；首末序号为1和605。BMI仅缺失1项，由训练折中位数填补，并附加缺失指示变量；孕周由“周+天”换算为连续周数。类别极不平衡：任一异常占11.1%，T21仅占2.1%，因此普通准确率会被多数类主导。")
    add_figure(doc, "raw_q4_label_distribution.png", "图1 女胎异常标签分布（记录级，标签可重叠）",
               "图1直接说明为什么不能以准确率作为主指标：若把所有记录都判为阴性，准确率仍会很高，却完全漏掉少数异常。故后文使用PR-AUC、灵敏度和F2作为主要判优依据。")
    pf.heading2(doc, "2.2 Z值与质量变量的初步关系")
    add_figure(doc, "raw_q4_zscore_distribution.png", "图2 不同AB标签下目标染色体绝对Z值分布",
               "传统方法常以固定Z值阈值筛查，但图2显示阳性与阴性记录的|Z|分布大量重叠，尤其T21阳性并未集中在|Z|≥3以上。单一阈值难以完整复现本附件的AB标签。")
    add_figure(doc, "raw_q4_quality_relationships.png", "图3 测序质量、孕妇特征与染色体Z值的相关结构",
               "图3显示Z13、Z18与ZX之间存在中等相关，GC与重复率也有关联。这支持同时利用多维质量变量，但相关并不等于因果，特征重要性只能解释模型使用情况。")

    pf.heading1(doc, "3 模型建立与算法设计")
    pf.heading2(doc, "3.1 为什么不用搜索树或Optimal Binning直接完成第四题")
    pf.body(doc, "搜索树和Optimal Binning适合寻找少量有序切点，例如前两问的BMI分组或推荐孕周；第四题的目标是从33个连续、比例和类别特征预测三个稀有标签。若枚举全部多维切点，搜索空间会组合爆炸，局部贪心还容易在小样本下产生脆弱规则。因此，本题不把“分组最优”误当成“分类最优”，而采用带正则的浅层提升树学习有限的非线性交互，并用严格的嵌套验证约束模型复杂度。")
    pf.heading2(doc, "3.2 多标签代价敏感XGBoost")
    pf.body(doc, "对每个染色体c∈{T13,T18,T21}建立独立二分类器。XGBoost以加性树逼近对数几率，能处理Z值、GC、读段质量和BMI之间的非线性关系[3]。为避免一位孕妇因重复检测较多而支配训练损失，先令每位孕妇的记录总权重相等；再按训练折阴阳记录比提高正类权重。")
    pf.equation(doc, r"\alpha_c=\frac{n_{c,-}}{n_{c,+}}")
    pf.body(doc, "候选模型仅设置三组浅树方案，内层训练以PR-AUC为主、F2为辅选择。这里的“浅”是有意的：147名孕妇不足以支持大深度搜索树，过多叶节点会记忆重复检测噪声。")
    pf.heading2(doc, "3.3 患者分组嵌套验证")
    pf.body(doc, "外层4折用于最终评价，内层3折用于特征填补、参数选择、概率校准与阈值选择。同一孕妇所有记录始终在同一折。自定义分折算法通过2000个确定性候选搜索，使四折的记录量和T13/T18/T21阳性患者分布尽量平衡；外层与内层的患者交集均为0。")
    add_figure(doc, "process_q4_fold_balance.png", "图4 外层四折患者规模与少数类覆盖",
               "图4表明每个外层测试折包含36—37名孕妇，并覆盖全部三个稀有标签。该图不是训练成绩，而是验证设计的质量检查。")
    pf.heading2(doc, "3.4 概率校准、任一异常与阈值")
    pf.body(doc, "超参数冻结后，以完整内层OOF原始概率拟合Platt sigmoid映射，再作用于对应外层测试概率；阳性不足8条或数值失败时回退为原始概率。三个染色体校准概率通过概率并集形成任一异常风险：")
    pf.equation(doc, r"p_i(any)=1-\prod_{c\in C}(1-p_{ic})")
    pf.body(doc, "概率并集与直接任一异常分类器只在每个外层训练集的OOF结果中比较。若并集PR-AUC或Brier比直接模型差超过0.02才回退，否则采用并集；四个外层训练折均选择并集。外层测试标签仅用于最后评价，未参与这一选择。")
    pf.body(doc, "阈值在训练OOF候选中选择：阳性不少于8条时，先保留灵敏度≥0.85的阈值，再最大化特异度，若并列则比较F2并取较高阈值。阈值不是固定0.5，因为稀有标签下0.5会造成严重漏检。")
    add_figure(doc, "process_q4_threshold_tradeoff.png", "图5 任一异常训练OOF阈值的灵敏度—特异度—F2权衡",
               "图5说明阈值越低，灵敏度越高而特异度越低。最终规则把漏检风险置于优先位置，因此选择点不一定使精确率最大。")
    pf.heading2(doc, "3.5 不确定带与质量复检")
    pf.body(doc, "若概率非常接近阈值，强制给出阳性或阴性会夸大模型确定性。对每个染色体，以训练OOF概率到阈值距离的15% lower经验分位数确定带宽，上限0.15；相同距离按原始序号稳定排序，单折最多标记floor(0.15n)条。")
    pf.equation(doc, r"\delta_c=\min(0.15,Q_{0.15}(|p_{ic}-t_c|))")
    pf.body(doc, "质量极端由总读段、唯一比对读段、比对率、重复率、过滤率和总GC六项判断：至少两项越过训练折1%—99%范围，或总/唯一读段低于0.5%分位点，即建议重抽复检。所有边界仅在训练折估计。")

    pf.heading1(doc, "4 评分指标及通俗解释")
    pf.body(doc, "本题阳性稀少，单看准确率会掩盖漏检。各指标回答的问题不同，必须联合解释。PR-AUC衡量从高风险到低风险排序时对少数阳性的发现能力；灵敏度回答“真实AB阳性中找出多少”；特异度回答“真实AB阴性中排除多少”；精确率回答“模型报阳中有多少确为AB阳性”；F2把灵敏度权重设为精确率的4倍；Brier衡量概率与0/1标签的平均平方误差，越小越好，但不同患病率标签之间不宜直接用Brier横比。")
    add_table(doc, [["指标", "范围/方向", "本题含义"],
                    ["PR-AUC", "0—1，越高越好", "不平衡数据下阳性排序质量"],
                    ["灵敏度", "0—1，越高漏检越少", "AB阳性被找出的比例"],
                    ["特异度", "0—1，越高误报越少", "AB阴性被正确排除的比例"],
                    ["精确率", "0—1，越高阳性更可信", "筛查阳性中AB阳性的比例"],
                    ["F2", "0—1，越高越好", "更重视漏检的综合分数"],
                    ["Brier", "0—1，越低越好", "概率预测的平均平方误差"],
                    ["Bootstrap 95%区间", "越窄越稳定", "以孕妇为单位重采样的波动范围"]],
              "表1 问题4评分指标及解释", [2.7, 3.5, 9.0])
    pf.equation(doc, r"F_2=\frac{5PR}{4P+R}")

    pf.heading1(doc, "5 验证结果与模型比较")
    pf.heading2(doc, "5.1 完全留出结果")
    result_rows = [["目标", "PR-AUC", "灵敏度", "特异度", "精确率", "F2", "Brier"]]
    for lab in ("T13", "T18", "T21", "any"):
        label_name = "任一异常" if lab == "any" else lab
        result_rows.append([label_name, fmt(b(lab, "pr_auc")[0]), fmt(b(lab, "sensitivity")[0]),
                            fmt(b(lab, "specificity")[0]), fmt(b(lab, "precision")[0]),
                            fmt(b(lab, "f2")[0]), fmt(b(lab, "brier")[0])])
    add_table(doc, result_rows, "表2 外层pooled OOF观测指标", [2.1, 2.1, 2.1, 2.1, 2.1, 2.1, 2.1])
    pf.body(doc, "T18表现相对最好，PR-AUC为0.412、F2为0.406；T13有一定排序能力但精确率低；T21的PR-AUC仅0.033，说明当前13条阳性不足以形成稳定区分。任一异常综合模型PR-AUC为0.446、灵敏度为0.836，但特异度0.409、精确率0.150，适合高敏感度分流，不适合直接确诊。")
    add_figure(doc, "result_q4_pr_curves.png", "图6 外层留出预测的PR曲线（图例AP为pooled OOF口径）",
               "图6显示T18和任一异常的曲线明显优于相应阳性率基准；T21曲线接近基准，提示其结论必须保守。")
    pf.heading2(doc, "5.2 概率校准与不确定性")
    add_figure(doc, "result_q4_calibration.png", "图7 外层留出概率校准曲线与Brier分数",
               "图7按五个等频概率箱比较平均预测概率与实际阳性率。少数类每箱阳性数有限，因此曲线波动不能解释为真实临床风险校准，只用于检查AB标签概率是否严重偏离。")
    ci_rows = [["目标", "灵敏度（95% CI）", "PR-AUC（95% CI）", "F2（95% CI）"]]
    for lab in ("T13", "T18", "T21", "any"):
        name = "任一异常" if lab == "any" else lab
        s, ap, f2 = b(lab, "sensitivity"), b(lab, "pr_auc"), b(lab, "f2")
        ci_rows.append([name, f"{s[0]:.3f} [{s[1]:.3f}, {s[2]:.3f}]",
                        f"{ap[0]:.3f} [{ap[1]:.3f}, {ap[2]:.3f}]",
                        f"{f2[0]:.3f} [{f2[1]:.3f}, {f2[2]:.3f}]"])
    add_table(doc, ci_rows, "表3 患者级2000次Bootstrap稳定性", [2.2, 4.3, 4.3, 4.3])
    pf.heading2(doc, "5.3 与传统Z规则的对照")
    pf.body(doc, "固定|Z|≥3在本附件上几乎不能复现AB标签：T13 pooled灵敏度约0.043，T18和T21均为0。训练折内调节Z阈值可提高灵敏度，却带来更低特异度；多变量模型在F2上整体更优。这并不否定医学Z值方法，而说明附件AB标签的形成机制与单一固定阈值并不一致。")
    add_figure(doc, "result_q4_metrics_comparison.png", "图8 多变量模型与两种Z值基线的四外折均值比较",
               "图8使用四个外层折指标的算术均值，口径不同于表2的pooled OOF；二者分别描述典型折表现和全部留出记录整体表现。固定Z规则的零高度柱已标注为0。")
    pf.heading2(doc, "5.4 模型使用了哪些变量")
    add_figure(doc, "result_q4_feature_importance.png", "图9 三个染色体模型的特征重要性汇总",
               "图9按三个标签×四个外层折共12个模型汇总重要性，误差线同时包含标签差异和折间差异。X染色体浓度、染色体GC、GC偏离、BMI和过滤比例较常被使用；重要性不代表因果效应，也不能据此提出医学干预。")

    pf.heading1(doc, "6 最终判定方法与建议表")
    pf.heading2(doc, "6.1 四级互斥分流规则")
    pf.body(doc, "为避免同一记录同时落入多个动作，采用明确优先级。先检查测序质量，再检查三个染色体任一概率是否处于不确定带；只有前两项均未触发时，才按任一异常模型阈值给出筛查阳性或阴性。T13/T18/T21单项概率用于阳性后的风险归因，不等同于确诊染色体类型。")
    action_order = ["质量异常：建议重抽复检", "阈值附近：建议复检", "筛查阳性：建议遗传咨询及诊断性检查", "筛查阴性：常规随访"]
    action_rows = [["优先级", "互斥动作", "记录数", "解释"]]
    explanations = ["六项质量规则触发，先保证检测可靠", "任一T13/T18/T21概率接近训练阈值", "任一异常超过训练折阈值，不作为确诊", "未触发前述规则，仍需常规产检"]
    for i, (name, exp) in enumerate(zip(action_order, explanations), 1):
        action_rows.append([i, name, int(actions.get(name, 0)), exp])
    add_table(doc, action_rows, "表4 最终建议表（互斥、按优先级分流）", [1.5, 5.6, 2.0, 6.2])
    pf.heading2(doc, "6.2 建议表的使用方式")
    add_bullets(doc, ["质量异常（17条）：优先重抽或复测，不解释当前概率。",
                      "阈值附近（190条）：模型无法稳定二分，复检优先于阳性/阴性结论。",
                      "筛查阳性（253条）：建议遗传咨询与诊断性检查；低精确率意味着不能直接宣告异常。",
                      "筛查阴性（145条）：按常规随访；模型仍存在11条AB阳性漏检，不能替代临床流程。",
                      "逐记录概率和动作已保存于results/q4_final_recommendations.csv，可按序号回查。"])

    pf.heading1(doc, "7 稳健性、优点与局限")
    pf.heading2(doc, "7.1 稳健性证据")
    add_bullets(doc, [f"患者隔离：4个外折和12个内折的患者交集最大值为{int(audit.patient_overlap.max())}。",
                      "划分稳定：每次分折从2000个确定性候选中选择平衡解，所有稀有标签在每折均有阳性患者。",
                      "参数隔离：填补、权重、超参数、Platt校准、阈值和任一异常来源都在外层训练集内冻结。",
                      "区间评估：以孕妇为单位进行2000次Bootstrap，避免把同一孕妇多条记录当成独立个体。",
                      "复现绑定：附件、代码、结果、图及版本信息写入results/复现清单.json。"])
    pf.heading2(doc, "7.2 方法优点")
    pf.body(doc, "本方案针对重复检测、小样本、多标签和类别不平衡同时设计：患者分组阻断泄漏；患者等权防止高频复测个体主导；浅层正则化提升树控制复杂度；训练折阈值明确体现漏检优先；质量与不确定带把模型无法可靠处理的记录转为复检，而不是制造过度自信。")
    pf.heading2(doc, "7.3 主要局限")
    add_bullets(doc, ["结局局限：AB是筛查标签，AE全为“是”，没有真实核型或出生结局。",
                      "样本局限：仅147名孕妇，T21阳性患者尤其少，置信区间较宽。",
                      "误报局限：任一异常精确率约0.150，阳性结果必须进入二级检查。",
                      "校准局限：内部Platt校准不能替代独立医院、平台和人群的外部校准。",
                      "解释局限：树重要性反映预测使用频率，不证明BMI、GC等变量导致异常。"])

    pf.heading1(doc, "8 结论")
    pf.body(doc, "第四题的合理完成方案不是继续扩大搜索树，而是把问题建模为患者分组、多标签、代价敏感的概率筛查。最终模型在无患者泄漏的外层留出预测上，任一异常PR-AUC为0.446、灵敏度为0.836、F2为0.436；固定|Z|≥3在附件AB标签上漏检严重。四个外层训练折均选择三个染色体概率并集，说明该组合方向在内部数据上具有一致性。")
    pf.body(doc, "然而，特异度0.409、精确率0.150以及T21极低PR-AUC表明不存在可直接临床应用的“稳定诊断方案”。可稳定交付的是一套经过验证的研究性分流流程：17条质量异常先复检，190条阈值附近复检，253条筛查阳性进入遗传咨询与诊断性检查，145条筛查阴性常规随访。若获得真实核型/出生结局或新增独立队列，应以本框架重新训练并完成外部验证。")

    pf.heading1(doc, "参考文献")
    refs = [
        "[1] Koumbaris G, Kypri E, Tsangaras K, et al. Cell-Free DNA Analysis of Targeted Genomic Regions in Maternal Plasma for Non-Invasive Prenatal Testing of Trisomy 21, Trisomy 18, Trisomy 13, and Fetal Sex. Clinical Chemistry, 2016, 62(6): 848-855. DOI: 10.1373/clinchem.2015.252502.",
        "[2] Lee J, Lee S M, Ahn J M, et al. Development and performance evaluation of an artificial intelligence algorithm using cell-free DNA fragment distance for non-invasive prenatal testing (aiD-NIPT). Frontiers in Genetics, 2022, 13: 999587. DOI: 10.3389/fgene.2022.999587.",
        "[3] Chen T, Guestrin C. XGBoost: A Scalable Tree Boosting System. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 2016: 785-794. DOI: 10.1145/2939672.2939785.",
    ]
    for ref in refs:
        p = pf.paragraph(doc, ref)
        p.paragraph_format.first_line_indent = Cm(-0.74); p.paragraph_format.left_indent = Cm(0.74)

    pf.heading1(doc, "附录A 复现文件说明")
    add_table(doc, [["文件", "作用"],
                    ["question4/solve_q4.py", "数据校验、特征、分折、权重、阈值与smoke测试"],
                    ["question4/run_q4_full.py", "完整外4内3嵌套验证、基线、Bootstrap和建议表"],
                    ["question4/plot_q4.py", "10族原始/过程/结果图及图表合同"],
                    ["results/q4_outer_predictions.csv", "每条记录的外层留出概率、判定与不确定标记"],
                    ["results/q4_final_recommendations.csv", "最终逐记录建议表"],
                    ["results/q4_bootstrap_summary.csv", "患者级2000次Bootstrap点估计和95%区间"],
                    ["results/复现清单.json", "输入、代码、结果、图、依赖版本与SHA-256绑定"]],
              "表5 问题4权威代码与结果文件", [5.8, 9.5])
    pf.body(doc, "建议复现顺序：先运行run_q4_full.py生成全量结果，再运行plot_q4.py生成图，最后运行build_manifest.py冻结当前哈希快照。全量运行使用固定随机种子20250824；若修改任何代码、结果、图或报告，应重新生成复现清单并重新执行相应质检。")

    issues = pf.validate_paper_structure(doc, contest="cumcm")
    errors = [x for x in issues if not x.startswith("预警：")]
    if errors:
        raise ValueError("报告结构校验失败：" + "；".join(errors))
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
